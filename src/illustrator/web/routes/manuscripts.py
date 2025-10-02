"""API routes for manuscript management."""

import base64
import json
import os
import uuid
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Iterable

from fastapi import APIRouter, HTTPException, status, BackgroundTasks
from fastapi.responses import JSONResponse

from illustrator.models import (
    ManuscriptMetadata,
    SavedManuscript,
    Chapter,
    ImageProvider,
    EmotionalMoment,
    EmotionalTone,
    IllustrationPrompt,
    LLMProvider,
)
from illustrator.web.models.web_models import (
    ManuscriptCreateRequest,
    ManuscriptResponse,
    DashboardStats,
    ErrorResponse,
    SuccessResponse,
    StyleConfigSaveRequest
)
from illustrator.prompt_engineering import (
    StyleTranslator,
    SceneComposition,
    CompositionType,
    LightingMood,
    PromptEngineer,
)
from illustrator.analysis import EmotionalAnalyzer
from illustrator.context import get_default_context
from illustrator.llm_factory import create_chat_model_from_context

router = APIRouter()

logger = logging.getLogger(__name__)

# Storage paths
SAVED_MANUSCRIPTS_DIR = Path("saved_manuscripts")
ILLUSTRATOR_OUTPUT_DIR = Path("illustrator_output")
GENERATED_IMAGES_DIR = ILLUSTRATOR_OUTPUT_DIR / "generated_images"
PREVIEW_IMAGES_DIR = GENERATED_IMAGES_DIR / "previews"

# Common image extensions we manage on disk
_IMAGE_FILE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"}

# Ensure directories exist
SAVED_MANUSCRIPTS_DIR.mkdir(exist_ok=True)
ILLUSTRATOR_OUTPUT_DIR.mkdir(exist_ok=True)
GENERATED_IMAGES_DIR.mkdir(exist_ok=True)
PREVIEW_IMAGES_DIR.mkdir(exist_ok=True)


def generate_manuscript_id() -> str:
    """Generate a unique manuscript ID."""
    return str(uuid.uuid4())


def get_manuscript_processing_status(manuscript_id: str) -> str:
    """Get the current processing status of a manuscript."""
    try:
        from illustrator.web.app import connection_manager
        if hasattr(connection_manager, 'sessions'):
            for session_id, session_data in connection_manager.sessions.items():
                if session_data.manuscript_id == manuscript_id:
                    if hasattr(session_data, 'status') and session_data.status:
                        return session_data.status.status
                    return "processing"
        return "draft"
    except:
        return "draft"


# Cache for manuscripts to avoid repeated file system scans
_manuscripts_cache = {}
_cache_timestamp = None

# Shared style translator instance for preview generation
_style_translator = StyleTranslator()

_ROUTES_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _ROUTES_DIR.parents[4]


def _normalize_style_modifiers(modifiers: Iterable[Any] | None) -> List[str]:
    """Convert style modifiers into a list of readable strings."""
    if modifiers is None:
        return []

    normalized: List[str] = []
    try:
        for modifier in modifiers:
            if isinstance(modifier, (list, tuple)):
                normalized.append(" ".join(str(part) for part in modifier if part))
            else:
                normalized.append(str(modifier))
    except TypeError:
        # Handle case where modifiers is not iterable
        logger.warning("Received non-iterable style modifiers: %r", modifiers)
        return []

    return [m for m in normalized if m]


def _resolve_style_config_path(style_config_path: str) -> Optional[Path]:
    """Resolve possible locations for a rich style configuration file."""
    candidate_paths = []
    raw_path = Path(style_config_path)

    if raw_path.is_absolute():
        candidate_paths.append(raw_path)
    else:
        candidate_paths.extend([
            raw_path,
            _PROJECT_ROOT / raw_path,
            SAVED_MANUSCRIPTS_DIR / raw_path,
            SAVED_MANUSCRIPTS_DIR / "style_configs" / raw_path,
        ])

    seen = set()
    for candidate in candidate_paths:
        resolved = candidate.resolve()
        if resolved in seen:
            continue
        seen.add(resolved)
        if resolved.exists():
            return resolved
    return None


def _load_rich_style_config(style_config_path: str) -> Optional[Dict[str, Any]]:
    """Load a detailed style configuration file if available."""
    resolved_path = _resolve_style_config_path(style_config_path)
    if not resolved_path:
        logger.debug("Rich style config not found for path %s", style_config_path)
        return None

    try:
        with open(resolved_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as exc:  # noqa: BLE001 - we want to log and continue
        logger.warning("Failed to load rich style config %s: %s", resolved_path, exc)
        return None

def invalidate_manuscripts_cache():
    """Invalidate the manuscripts cache to force a refresh."""
    global _manuscripts_cache, _cache_timestamp
    _manuscripts_cache.clear()
    _cache_timestamp = None

def get_saved_manuscripts() -> List[SavedManuscript]:
    """Load all saved manuscripts from disk with caching."""
    global _manuscripts_cache, _cache_timestamp

    if not SAVED_MANUSCRIPTS_DIR.exists():
        return []

    # Check if cache is still valid (cache for 30 seconds)
    import time
    current_time = time.time()

    if _cache_timestamp and current_time - _cache_timestamp < 30:
        return [manuscript for manuscript, _ in _manuscripts_cache.values()]

    manuscripts = []
    new_cache = {}

    for file_path in SAVED_MANUSCRIPTS_DIR.glob("*.json"):
        try:
            # Check file modification time to see if we need to reload
            file_mtime = file_path.stat().st_mtime
            cache_key = str(file_path)

            # Use cached version if file hasn't changed
            if cache_key in _manuscripts_cache:
                cached_manuscript, cached_mtime = _manuscripts_cache[cache_key]
                if file_mtime <= cached_mtime:
                    new_cache[cache_key] = (cached_manuscript, cached_mtime)
                    manuscripts.append(cached_manuscript)
                    continue

            # Load from file
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Import validation helpers
            from illustrator.utils.validation_helpers import ensure_chapter_required_fields
            data = ensure_chapter_required_fields(data)
                
            manuscript = SavedManuscript(**data)
            new_cache[cache_key] = (manuscript, file_mtime)
            manuscripts.append(manuscript)

        except Exception as e:
            print(f"Error loading manuscript {file_path}: {e}")
            continue

    # Update cache
    _manuscripts_cache = new_cache
    _cache_timestamp = current_time

    # Sort by saved date (newest first)
    manuscripts.sort(key=lambda m: m.saved_at, reverse=True)
    return manuscripts


def save_manuscript_to_disk(manuscript: SavedManuscript) -> Path:
    """Save a manuscript to disk."""
    # Generate safe filename
    safe_title = "".join(c for c in manuscript.metadata.title if c.isalnum() or c in (' ', '-', '_')).strip()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{safe_title}_{timestamp}.json"
    file_path = SAVED_MANUSCRIPTS_DIR / filename

    # Update file path in manuscript
    manuscript.file_path = str(file_path)

    # Prepare data for saving, ensuring all required fields
    data = manuscript.model_dump()
    
    # Import validation helpers
    from illustrator.utils.validation_helpers import validate_manuscript_before_save
    validated_data = validate_manuscript_before_save(data)

    # Save to file
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(validated_data, f, indent=2, ensure_ascii=False)

    return file_path


def _normalized_title(metadata: ManuscriptMetadata) -> str:
    """Create a filesystem-safe title for locating manuscript assets."""
    safe_title = "".join(
        c for c in metadata.title if c.isalnum() or c in (" ", "-", "_")
    ).strip()
    return safe_title.replace(" ", "_")


def _filesystem_image_dirs_for_manuscript(manuscript_id: str) -> List[Path]:
    """Determine likely directories that contain generated images for a manuscript."""

    candidate_dirs: List[Path] = []

    # Prefer manuscript-specific directories inside the shared generated_images folder.
    manuscript_dir = GENERATED_IMAGES_DIR / manuscript_id
    if manuscript_dir.exists():
        candidate_dirs.append(manuscript_dir)

    # Check for title-based directories used by CLI flows.
    try:
        for manuscript in get_saved_manuscripts():
            generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
            if generated_id != manuscript_id:
                continue

            title_dir = ILLUSTRATOR_OUTPUT_DIR / _normalized_title(manuscript.metadata)
            if title_dir.exists():
                candidate_dirs.append(title_dir)

            nested_dir = title_dir / "generated_images"
            if nested_dir.exists():
                candidate_dirs.append(nested_dir)
            break
    except Exception:
        # If discovery fails we simply fall back to shared directories below.
        pass

    # If nothing specific was found, include the shared directory as a fallback.
    if not candidate_dirs and GENERATED_IMAGES_DIR.exists():
        candidate_dirs.append(GENERATED_IMAGES_DIR)

    # Deduplicate while keeping order.
    seen: set[Path] = set()
    unique_dirs: List[Path] = []
    for directory in candidate_dirs:
        try:
            resolved = directory.resolve()
        except (OSError, RuntimeError):
            continue
        if resolved in seen:
            continue
        seen.add(resolved)
        unique_dirs.append(directory)

    return unique_dirs


def filesystem_count_generated_images(manuscript_id: str, chapter_number: int | None = None) -> int:
    """Count generated images on disk, optionally filtered by chapter number."""

    directories = _filesystem_image_dirs_for_manuscript(manuscript_id)
    if not directories:
        return 0

    try:
        previews_dir_resolved = PREVIEW_IMAGES_DIR.resolve()
    except (OSError, RuntimeError):
        previews_dir_resolved = None

    seen_files: set[Path] = set()
    chapter_tokens: tuple[str, ...] | None = None
    if chapter_number is not None:
        chapter_tokens = (
            f"chapter_{chapter_number}_",
            f"chapter_{chapter_number:02d}_",
        )

    total = 0
    for directory in directories:
        if not directory.exists():
            continue
        try:
            iterator = directory.rglob("*")
        except (OSError, RuntimeError):
            continue

        for path in iterator:
            if not path.is_file():
                continue
            if path.suffix.lower() not in _IMAGE_FILE_EXTENSIONS:
                continue
            if chapter_tokens and not any(token in path.stem for token in chapter_tokens):
                continue

            try:
                resolved = path.resolve()
            except (OSError, RuntimeError):
                continue

            if (
                previews_dir_resolved
                and (
                    resolved == previews_dir_resolved
                    or previews_dir_resolved in resolved.parents
                )
            ):
                continue

            if resolved in seen_files:
                continue

            seen_files.add(resolved)
            total += 1

    return total


def filesystem_delete_generated_images(manuscript_id: str) -> int:
    """Delete generated images for a manuscript from the filesystem."""

    directories = _filesystem_image_dirs_for_manuscript(manuscript_id)
    if not directories:
        return 0

    removed = 0
    for directory in directories:
        if not directory.exists():
            continue

        try:
            iterator = directory.rglob("*")
        except (OSError, RuntimeError):
            continue

        for path in iterator:
            if not path.is_file():
                continue
            if path.suffix.lower() not in _IMAGE_FILE_EXTENSIONS:
                continue
            try:
                path.unlink()
                removed += 1
            except OSError:
                # Ignore deletion issues so we can report partial cleanup.
                continue

        # Attempt to prune empty manuscript-specific directories while leaving shared ones intact.
        if directory not in (GENERATED_IMAGES_DIR, ILLUSTRATOR_OUTPUT_DIR):
            try:
                next(directory.iterdir())
            except StopIteration:
                try:
                    directory.rmdir()
                except OSError:
                    pass
            except (OSError, RuntimeError):
                pass

    # Additionally clear previews if they exist and we managed to remove anything else.
    if removed and PREVIEW_IMAGES_DIR.exists():
        try:
            for preview_file in PREVIEW_IMAGES_DIR.iterdir():
                if not preview_file.is_file():
                    continue
                if preview_file.suffix.lower() not in _IMAGE_FILE_EXTENSIONS:
                    continue
                if not preview_file.name.startswith(manuscript_id):
                    continue
                try:
                    preview_file.unlink()
                except OSError:
                    continue
        except (OSError, RuntimeError):
            pass

    return removed


def count_generated_images(manuscript_id: str) -> int:
    """Count generated images for a manuscript.

    Attempts to use the database when available, falling back to a coarse
    filesystem count when the database layer isn't configured.
    """

    db_count = 0
    illustration_service = None

    try:
        from illustrator.services.illustration_service import IllustrationService

        illustration_service = IllustrationService()
        db_count = len(illustration_service.get_illustrations_by_manuscript(manuscript_id))
    except Exception:
        db_count = 0
    finally:
        if illustration_service is not None:
            try:
                illustration_service.close()
            except Exception:
                pass

    if db_count:
        return db_count

    return filesystem_count_generated_images(manuscript_id)


@router.get("/stats")
async def get_dashboard_stats() -> DashboardStats:
    """Get dashboard statistics."""
    logger.debug("API CALL: GET /stats")
    manuscripts = get_saved_manuscripts()

    total_chapters = sum(len(m.chapters) for m in manuscripts)
    total_images = 0
    for manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        total_images += count_generated_images(generated_id)
    # Track active processing sessions by checking connection manager
    from illustrator.web.models.web_models import ConnectionManager
    try:
        # Check for active processing sessions
        processing_count = 0
        from illustrator.web.app import connection_manager
        if hasattr(connection_manager, 'sessions'):
            processing_count = len(connection_manager.sessions)
    except:
        processing_count = 0

    # Get recent manuscripts (last 10)
    recent_manuscripts = []
    for manuscript in manuscripts[:10]:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        recent_manuscripts.append(ManuscriptResponse(
            id=generated_id,
            metadata=manuscript.metadata,
            chapters=manuscript.chapters,
            total_images=count_generated_images(generated_id),
            processing_status=get_manuscript_processing_status(generated_id),
            created_at=manuscript.metadata.created_at,
            updated_at=manuscript.saved_at
        ))

    return DashboardStats(
        total_manuscripts=len(manuscripts),
        total_chapters=total_chapters,
        total_images=total_images,
        recent_manuscripts=recent_manuscripts,
        processing_count=processing_count
    )


@router.get("/")
async def list_manuscripts() -> List[ManuscriptResponse]:
    """List all manuscripts."""
    logger.debug("API CALL: GET / (list_manuscripts)")
    manuscripts = get_saved_manuscripts()

    response = []
    for manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        response.append(ManuscriptResponse(
            id=generated_id,
            metadata=manuscript.metadata,
            chapters=manuscript.chapters,
            total_images=count_generated_images(generated_id),
            processing_status=get_manuscript_processing_status(generated_id),
            created_at=manuscript.metadata.created_at,
            updated_at=manuscript.saved_at
        ))

    return response


@router.get("/{manuscript_id}")
async def get_manuscript(manuscript_id: str) -> ManuscriptResponse:
    """Get a specific manuscript by ID."""
    logger.debug(f"API CALL: GET /{{manuscript_id}} (get_manuscript) with manuscript_id={manuscript_id}")
    manuscripts = get_saved_manuscripts()

    # Find manuscript by generated ID
    for manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        if generated_id == manuscript_id:
            return ManuscriptResponse(
                id=manuscript_id,
                metadata=manuscript.metadata,
                chapters=manuscript.chapters,
                total_images=count_generated_images(manuscript_id),
                processing_status=get_manuscript_processing_status(generated_id),
                created_at=manuscript.metadata.created_at,
                updated_at=manuscript.saved_at
            )

    raise HTTPException(
        status_code=status.HTTP_404_NOT_FOUND,
        detail="Manuscript not found"
    )


@router.post("/")
async def create_manuscript(request: ManuscriptCreateRequest) -> ManuscriptResponse:
    """Create a new manuscript."""
    logger.debug(f"API CALL: POST / (create_manuscript) with title={request.title}")
    # Create manuscript metadata
    metadata = ManuscriptMetadata(
        title=request.title,
        author=request.author,
        genre=request.genre,
        total_chapters=0,
        created_at=datetime.now().isoformat()
    )

    # Create saved manuscript
    saved_manuscript = SavedManuscript(
        metadata=metadata,
        chapters=[],
        saved_at=datetime.now().isoformat(),
        file_path=""  # Will be set when saving
    )

    # Save to disk
    file_path = save_manuscript_to_disk(saved_manuscript)
    manuscript_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, str(file_path)))

    # Invalidate cache to ensure new manuscript appears in the dashboard
    invalidate_manuscripts_cache()

    return ManuscriptResponse(
        id=manuscript_id,
        metadata=metadata,
        chapters=[],
        total_images=0,
        processing_status="draft",
        created_at=metadata.created_at,
        updated_at=saved_manuscript.saved_at
    )


@router.put("/{manuscript_id}")
async def update_manuscript(
    manuscript_id: str,
    request: ManuscriptCreateRequest
) -> ManuscriptResponse:
    """Update an existing manuscript."""
    logger.debug(f"API CALL: PUT /{{manuscript_id}} (update_manuscript) with manuscript_id={manuscript_id}")
    manuscripts = get_saved_manuscripts()

    # Find and update manuscript
    for manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        if generated_id == manuscript_id:
            # Update metadata
            manuscript.metadata.title = request.title
            manuscript.metadata.author = request.author
            manuscript.metadata.genre = request.genre
            manuscript.saved_at = datetime.now().isoformat()

            # Save updated manuscript
            save_manuscript_to_disk(manuscript)

            # Invalidate cache to ensure updated manuscript appears in the dashboard
            invalidate_manuscripts_cache()

            return ManuscriptResponse(
                id=manuscript_id,
                metadata=manuscript.metadata,
                chapters=manuscript.chapters,
                total_images=count_generated_images(manuscript_id),
                processing_status="draft",
                created_at=manuscript.metadata.created_at,
                updated_at=manuscript.saved_at
            )

    raise HTTPException(
        status_code=status.HTTP_404_NOT_FOUND,
        detail="Manuscript not found"
    )


@router.post("/{manuscript_id}/style")
async def save_style_config(
    manuscript_id: str,
    request: StyleConfigSaveRequest
) -> SuccessResponse:
    """Save style configuration for a manuscript."""
    logger.debug(f"API CALL: POST /{{manuscript_id}}/style (save_style_config) with manuscript_id={manuscript_id}")
    # Verify the manuscript exists first
    manuscripts = get_saved_manuscripts()

    manuscript_found = None
    for manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        if generated_id == manuscript_id:
            manuscript_found = manuscript
            break

    if not manuscript_found:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Manuscript not found"
        )

    try:
        # Create a style config directory for this manuscript if it doesn't exist
        style_config_dir = SAVED_MANUSCRIPTS_DIR / "style_configs"
        style_config_dir.mkdir(exist_ok=True)

        # Save the style configuration
        style_config_file = style_config_dir / f"{manuscript_id}_style.json"

        with open(style_config_file, 'w', encoding='utf-8') as f:
            json.dump(
                request.style_config.model_dump(exclude={'replicate_model'}),
                f,
                indent=2,
                ensure_ascii=False,
            )

        return SuccessResponse(
            message="Style configuration saved successfully",
            data={"style_config_path": str(style_config_file)}
        )

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error saving style configuration: {str(e)}"
        )


@router.get("/{manuscript_id}/style")
async def get_style_config(manuscript_id: str) -> Dict[str, Any]:
    """Get saved style configuration for a manuscript."""
    logger.debug(f"API CALL: GET /{{manuscript_id}}/style (get_style_config) with manuscript_id={manuscript_id}")
    # Verify the manuscript exists first
    manuscripts = get_saved_manuscripts()

    for manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        if generated_id == manuscript_id:
            break
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Manuscript not found"
        )

    try:
        # Check if style configuration exists
        style_config_dir = SAVED_MANUSCRIPTS_DIR / "style_configs"
        style_config_file = style_config_dir / f"{manuscript_id}_style.json"

        if not style_config_file.exists():
            return {"style_config": None}

        with open(style_config_file, 'r', encoding='utf-8') as f:
            style_config_data = json.load(f)

        return {"style_config": style_config_data}

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error loading style configuration: {str(e)}"
        )


@router.post("/{manuscript_id}/style/preview")
async def preview_style_image(
    manuscript_id: str,
    request: StyleConfigSaveRequest,
    background_tasks: BackgroundTasks
) -> Dict[str, Any]:
    """Generate a preview image using the style configuration."""
    logger.debug(f"API CALL: POST /{{manuscript_id}}/style/preview (preview_style_image) with manuscript_id={manuscript_id}")
    from illustrator.providers import get_image_provider
    from illustrator.models import StyleConfig

    # Verify the manuscript exists
    manuscripts = get_saved_manuscripts()
    for manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        if generated_id == manuscript_id:
            break
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Manuscript not found"
        )

    try:
        logger.debug("Starting preview generation for manuscript %s", manuscript_id)
        # Create style config from request
        style_config = StyleConfig(**request.style_config.model_dump(exclude={'replicate_model'}))
        logger.debug("Style config created: provider=%s", style_config.image_provider)

        # Get context early to ensure GCP Project ID is available for all providers
        context = get_default_context()
        logger.debug("Initial context LLM provider: %s", context.llm_provider)
        logger.debug("Initial context GCP Project ID: %s", getattr(context, 'gcp_project_id', None))
        
        # COMPREHENSIVE DEBUGGING: Check all environment variables
        logger.debug("=== ENVIRONMENT VARIABLE DEBUG ===")
        logger.debug("GOOGLE_PROJECT_ID: %s", os.getenv('GOOGLE_PROJECT_ID'))
        logger.debug("GCP_PROJECT_ID: %s", os.getenv('GCP_PROJECT_ID'))
        logger.debug("LLM_PROVIDER: %s", os.getenv('LLM_PROVIDER'))
        logger.debug("DEFAULT_LLM_PROVIDER: %s", os.getenv('DEFAULT_LLM_PROVIDER'))
        logger.debug("ANTHROPIC_API_KEY present: %s", bool(os.getenv('ANTHROPIC_API_KEY')))
        logger.debug("=== END ENVIRONMENT DEBUG ===")
        
        # Ensure user API configuration overrides defaults
        # Always check for GCP Project ID regardless of detected provider
        gcp_project_id = os.getenv('GOOGLE_PROJECT_ID') or os.getenv('GCP_PROJECT_ID')
        logger.debug("Environment GCP Project ID: %s", gcp_project_id)
        
        # Check if user selected Anthropic Vertex and has GCP project ID
        if context.llm_provider == LLMProvider.ANTHROPIC_VERTEX:
            logger.debug("Detected Anthropic Vertex provider, checking GCP Project ID...")
            
            # Check if the GCP Project ID is a placeholder or invalid
            is_placeholder = (
                not gcp_project_id or 
                gcp_project_id in ['your-google-project-id', 'your-gcp-project-id', 'placeholder'] or
                gcp_project_id.startswith('your-') or
                len(gcp_project_id.strip()) == 0
            )
            
            if is_placeholder:
                logger.error("GCP Project ID is missing or appears to be a placeholder: %s", gcp_project_id)
                
                # Check what's in the .env file to provide better guidance
                env_file_value = os.getenv('GOOGLE_PROJECT_ID')
                logger.debug("GOOGLE_PROJECT_ID from environment: %s", env_file_value)
                
                error_msg = (
                    "GCP Project ID is required for Anthropic Vertex provider.\n\n"
                    "🔧 REQUIRED STEPS:\n"
                    "1. Click the API Configuration button (⚙️ gear icon) in the top navigation\n"
                    "2. Enter your real Google Cloud Project ID (not 'your-google-project-id')\n"
                    "3. Select 'Anthropic Vertex' as LLM Provider\n"
                    "4. Enter your Anthropic API key if needed\n"
                    "5. Click 'Save Configuration'\n\n"
                    f"💡 Current value '{gcp_project_id}' appears to be a placeholder.\n"
                    f"Your .env file shows: GOOGLE_PROJECT_ID={env_file_value}\n\n"
                    "After saving, the API Configuration will override your .env file settings."
                )
                
                raise ValueError(error_msg)
            
            context.gcp_project_id = gcp_project_id
            logger.debug("Context GCP Project ID updated to: %s", context.gcp_project_id)
        else:
            logger.debug("Not using Anthropic Vertex provider (provider: %s), but setting GCP Project ID anyway", context.llm_provider)
            if gcp_project_id:
                context.gcp_project_id = gcp_project_id
                logger.debug("Context GCP Project ID set to: %s", context.gcp_project_id)

        # Get the image provider
        gcp_project_id_to_pass = getattr(context, 'gcp_project_id', None)
        anthropic_api_key_to_pass = getattr(context, 'anthropic_api_key', None)
        
        logger.debug("=== PROVIDER CREATION DEBUG ===")
        logger.debug("Image provider type: %s", style_config.image_provider)
        logger.debug("Passing GCP Project ID to provider: %s", gcp_project_id_to_pass)
        logger.debug("Passing LLM provider to provider: %s", context.llm_provider)
        logger.debug("Passing Anthropic API key present: %s", bool(anthropic_api_key_to_pass))
        logger.debug("Context has anthropic_api_key attr: %s", hasattr(context, 'anthropic_api_key'))
        logger.debug("Context attributes: %s", [attr for attr in dir(context) if not attr.startswith('_')])
        logger.debug("=== END PROVIDER DEBUG ===")
        
        provider = get_image_provider(
            style_config.image_provider,
            huggingface_image_model=style_config.huggingface_model_id,
            huggingface_image_provider=style_config.huggingface_provider,
            # Pass the GCP Project ID and LLM provider explicitly to ensure they're available
            gcp_project_id=gcp_project_id_to_pass,
            llm_provider=context.llm_provider,
            anthropic_api_key=anthropic_api_key_to_pass,
        )
        logger.debug("Image provider created successfully: %s", type(provider))

        # Merge base config with any linked rich configuration for accurate previews
        rich_config = None
        if style_config.style_config_path:
            rich_config = _load_rich_style_config(style_config.style_config_path)

        style_preferences = style_config.model_dump()
        style_preferences["image_provider"] = style_config.image_provider.value
        if rich_config:
            style_preferences = {**rich_config, **style_preferences}

        # Create a representative scene composition to translate style preferences
        scene_composition = SceneComposition(
            composition_type=CompositionType.MEDIUM_SHOT,
            focal_point="primary narrative subjects",
            background_elements=["supporting environment"],
            foreground_elements=["key characters"],
            lighting_mood=LightingMood.NATURAL,
            atmosphere="balanced preview for configured illustration style",
            color_palette_suggestion=style_config.color_palette or "cohesive tones",
            emotional_weight=0.45,
        )

        logger.debug("About to translate style config")
        try:
            style_translation = _style_translator.translate_style_config(
                style_preferences,
                style_config.image_provider,
                scene_composition
            )
            logger.debug("Style translation completed successfully")
            logger.debug("Style translation keys: %s", list(style_translation.keys()) if style_translation else "None")
            if style_translation and 'style_modifiers' in style_translation:
                logger.debug("Style modifiers type: %s, value: %r", type(style_translation['style_modifiers']), style_translation['style_modifiers'])
        except Exception as e:
            logger.error("Style translation failed: %s", str(e), exc_info=True)
            raise

        style_modifiers = _normalize_style_modifiers(style_translation.get("style_modifiers", []))
        is_flux_family = style_config.image_provider in {
            ImageProvider.FLUX,
            ImageProvider.FLUX_DEV_VERTEX,
            ImageProvider.SEEDREAM,
            ImageProvider.HUGGINGFACE,
        }

        if is_flux_family and len(style_modifiers) > 6:
            style_modifiers = style_modifiers[:6]
        provider_opts = style_translation.get("provider_optimizations", {}) or {}

        # Flux prompts are sensitive to length; trim verbose quality modifiers
        if is_flux_family and provider_opts.get("quality_modifiers"):
            provider_opts = {
                **provider_opts,
                "quality_modifiers": provider_opts["quality_modifiers"][:2],
            }

        technical_params = dict(style_translation.get("technical_params", {}) or {})
        technical_params.update(provider_opts.get("technical_adjustments", {}) or {})
        if style_config.huggingface_model_id:
            technical_params.setdefault('model_id', style_config.huggingface_model_id)
        if style_config.huggingface_provider:
            technical_params.setdefault('provider', style_config.huggingface_provider)

        negative_prompt_items = []
        negative_prompt_data = style_translation.get("negative_prompt", []) or []
        if negative_prompt_data:
            try:
                negative_prompt_items = [str(item) for item in negative_prompt_data]
            except TypeError:
                logger.warning("Received non-iterable negative_prompt data: %r", negative_prompt_data)
                negative_prompt_items = []

        # Build a baseline preview prompt from style settings
        prompt_sections: List[str] = []
        base_description = f"Concept illustration in {style_config.art_style}"
        if style_config.color_palette:
            base_description += f" using the {style_config.color_palette} palette"
        if style_config.artistic_influences:
            base_description += f" inspired by {style_config.artistic_influences}"
        prompt_sections.append(base_description)

        if style_modifiers:
            prompt_sections.append("featuring " + ", ".join(style_modifiers))
        if provider_opts.get("style_emphasis"):
            prompt_sections.append(provider_opts["style_emphasis"])
        if (
            provider_opts.get("quality_modifiers")
            and not is_flux_family
        ):
            prompt_sections.append(
                "quality focus: " + ", ".join(provider_opts["quality_modifiers"])
            )

        preview_prompt_text = ". ".join(section for section in prompt_sections if section)
        if not preview_prompt_text:
            preview_prompt_text = "Concept illustration showcasing the configured style settings"
        if not preview_prompt_text.endswith('.'):
            preview_prompt_text += '.'

        illustration_prompt = IllustrationPrompt(
            provider=style_config.image_provider,
            prompt=preview_prompt_text,
            style_modifiers=style_modifiers,
            negative_prompt=", ".join(negative_prompt_items) if negative_prompt_items else None,
            technical_params=technical_params
        )

        excerpt_text = (request.manuscript_excerpt or "").strip()
        logger.debug("Excerpt text provided: %s", bool(excerpt_text))
        logger.debug("Excerpt text length: %d", len(excerpt_text) if excerpt_text else 0)

        if excerpt_text:
            try:
                logger.debug("Starting excerpt analysis for preview generation")
                
                context.image_provider = style_config.image_provider
                context.default_art_style = style_config.art_style or context.default_art_style
                context.color_palette = style_config.color_palette or context.color_palette
                context.artistic_influences = style_config.artistic_influences or context.artistic_influences

                logger.debug("Creating LLM for excerpt analysis")
                logger.debug("Context LLM provider: %s", context.llm_provider)
                logger.debug("Context model: %s", context.model)
                logger.debug("Context has anthropic_api_key: %s", bool(getattr(context, 'anthropic_api_key', None)))
                logger.debug("Context has huggingface_api_key: %s", bool(getattr(context, 'huggingface_api_key', None)))

                llm = create_chat_model_from_context(context)
                logger.debug("LLM created successfully: %s", type(llm))
                emotional_analyzer = EmotionalAnalyzer(llm)

                excerpt_id = f"preview-{uuid.uuid4()}"
                chapter_excerpt = Chapter(
                    id=excerpt_id,
                    title="Preview Excerpt",
                    content=excerpt_text,
                    number=1,
                    word_count=len(excerpt_text.split()) or 1,
                    summary=f"Preview excerpt for style testing: {excerpt_text[:50]}{'...' if len(excerpt_text) > 50 else ''}",
                )

                logger.debug("Running emotional analysis on excerpt")
                analysis_result = await emotional_analyzer.analyze_chapter(
                    chapter_excerpt,
                    max_moments=1,
                    min_intensity=0.3,
                )
                logger.debug("Emotional analysis completed. Found %d moments", len(analysis_result) if analysis_result else 0)

                preview_moment = analysis_result[0] if analysis_result else None
                if not preview_moment:
                    logger.debug("No emotional moments found, creating fallback moment")
                    preview_moment = EmotionalMoment(
                        text_excerpt=excerpt_text[:240],
                        start_position=0,
                        end_position=min(len(excerpt_text), 240),
                        emotional_tones=[EmotionalTone.NEUTRAL],
                        intensity_score=0.5,
                        context=excerpt_text[:240],
                    )
                else:
                    logger.debug("Using emotional moment: %s", preview_moment.text_excerpt[:100] + "..." if len(preview_moment.text_excerpt) > 100 else preview_moment.text_excerpt)

                logger.debug("Creating prompt engineer and generating enhanced prompt")
                prompt_engineer = PromptEngineer(llm)
                engineered_prompt = await prompt_engineer.engineer_prompt(
                    emotional_moment=preview_moment,
                    provider=style_config.image_provider,
                    style_preferences=style_preferences,
                    chapter_context=chapter_excerpt,
                )

                style_modifiers = _normalize_style_modifiers(engineered_prompt.style_modifiers)
                preview_prompt_text = engineered_prompt.prompt
                technical_params = engineered_prompt.technical_params or technical_params

                logger.debug("Enhanced prompt generated: %s", preview_prompt_text[:200] + "..." if len(preview_prompt_text) > 200 else preview_prompt_text)
                logger.debug("Style modifiers: %r", style_modifiers[:3] if style_modifiers else [])

                illustration_prompt = IllustrationPrompt(
                    provider=style_config.image_provider,
                    prompt=preview_prompt_text,
                    style_modifiers=style_modifiers,
                    negative_prompt=engineered_prompt.negative_prompt,
                    technical_params=technical_params,
                )

            except Exception as excerpt_error:
                logger.error(
                    "Excerpt-aware preview generation failed: %s", excerpt_error,
                    exc_info=True,
                )
                logger.debug("Falling back to generic prompt due to analysis failure")

        # Generate the image
        logger.debug("About to generate image with provider %s", style_config.image_provider)
        logger.debug("Illustration prompt: prompt=%s, style_modifiers=%r", illustration_prompt.prompt[:100] + "..." if len(illustration_prompt.prompt) > 100 else illustration_prompt.prompt, illustration_prompt.style_modifiers)
        try:
            result = await provider.generate_image(illustration_prompt)
            logger.debug("Image generation completed successfully")
        except Exception as e:
            logger.error("Image generation failed: %s", str(e), exc_info=True)
            raise

        # Extract image URL from result
        metadata: Dict[str, Any] = {}
        if result.get('success'):
            if result.get('image_data'):
                # Persist preview image to disk so the browser can load it reliably
                image_bytes = base64.b64decode(result['image_data'])
                preview_filename = f"{manuscript_id}_style_preview_{uuid.uuid4().hex}.png"
                preview_path = PREVIEW_IMAGES_DIR / preview_filename

                preview_path.parent.mkdir(parents=True, exist_ok=True)

                with open(preview_path, 'wb') as preview_file:
                    preview_file.write(image_bytes)

                image_url = f"/generated/previews/{preview_filename}"
            else:
                # Use direct URL if available
                image_url = result.get('image_url', '')

            metadata = result.get('metadata', {}) or {}
            metadata.setdefault('prompt', illustration_prompt.prompt)
            metadata.setdefault('style_modifiers', illustration_prompt.style_modifiers)
            metadata.setdefault('technical_params', illustration_prompt.technical_params)
            metadata.setdefault('excerpt_used', bool(excerpt_text))
        else:
            status_code = result.get('status_code') or status.HTTP_502_BAD_GATEWAY
            error_message = result.get('error', 'Image generation failed')
            provider_label = style_config.image_provider.name.replace('_', ' ').title()
            raise HTTPException(
                status_code=status_code,
                detail=f"{provider_label} preview failed: {error_message}"
            )

        style_summary: Dict[str, Any] = {
            "provider": style_config.image_provider.value,
            "art_style": style_config.art_style,
            "color_palette": style_config.color_palette,
            "artistic_influences": style_config.artistic_influences,
            "style_config_path": style_config.style_config_path,
            "style_modifiers": style_modifiers,
        }

        if excerpt_text:
            excerpt_preview = " ".join(excerpt_text.split())
            if len(excerpt_preview) > 220:
                excerpt_preview = f"{excerpt_preview[:217]}..."
            style_summary["excerpt_preview"] = excerpt_preview
            style_summary["excerpt_used"] = True
        else:
            style_summary["excerpt_used"] = False

        if style_config.huggingface_model_id:
            style_summary["huggingface_model_id"] = style_config.huggingface_model_id
        if style_config.huggingface_provider:
            style_summary["huggingface_provider"] = style_config.huggingface_provider

        if rich_config and rich_config.get("style_name"):
            style_summary["style_name"] = rich_config["style_name"]
        if provider_opts.get("style_emphasis"):
            style_summary["style_emphasis"] = provider_opts["style_emphasis"]
        if provider_opts.get("quality_modifiers") and not is_flux_family:
            style_summary["quality_modifiers"] = provider_opts["quality_modifiers"]
        effective_prompt = metadata.get('prompt')
        if effective_prompt:
            style_summary["effective_prompt"] = effective_prompt
        if metadata.get('technical_params'):
            style_summary["technical_params_applied"] = metadata['technical_params']
        if metadata.get('style_modifiers'):
            style_summary["applied_style_modifiers"] = metadata['style_modifiers']

        if is_flux_family:
            effective_prompt = style_summary.get('effective_prompt', preview_prompt_text)
            if metadata.get('prompt_truncated'):
                style_summary["prompt_truncated"] = True
            preview_prompt_text = effective_prompt
        elif effective_prompt:
            preview_prompt_text = effective_prompt

        return {
            "image_url": image_url,
            "preview_prompt": preview_prompt_text,
            "style_summary": style_summary
        }

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating preview image: {str(e)}"
        )


@router.get("/{manuscript_id}/images")
async def list_manuscript_images(manuscript_id: str) -> Dict[str, Any]:
    """List all generated images for a manuscript from the database."""
    logger.debug(f"API CALL: GET /{{manuscript_id}}/images (list_manuscript_images) with manuscript_id={manuscript_id}")
    try:
        from illustrator.services.illustration_service import IllustrationService

        # Initialize illustration service
        illustration_service = IllustrationService()

        try:
            # Get illustrations from database
            illustrations = illustration_service.get_illustrations_by_manuscript(manuscript_id)

            # Convert to API format
            images = []
            for illustration in illustrations:
                chapter_number = illustration.get("chapter_number")
                chapter_title = illustration.get("chapter_title")
                if chapter_number is not None:
                    base_chapter = f"Chapter {chapter_number}"
                    chapter_info = f"{base_chapter}: {chapter_title}" if chapter_title else base_chapter
                else:
                    chapter_info = "Unknown Chapter"

                created_at = illustration.get("created_at")
                if isinstance(created_at, datetime):
                    created_at_str = created_at.isoformat()
                else:
                    created_at_str = created_at

                image_data = {
                    "id": str(illustration.get("id")),
                    "url": illustration.get("web_url"),
                    "filename": illustration.get("filename"),
                    "title": illustration.get("title") or f"{chapter_info} - Scene {illustration.get('scene_number')}",
                    "description": illustration.get("description") or (
                        f"Generated illustration for {chapter_info}, Scene {illustration.get('scene_number')}"
                    ),
                    "chapter": chapter_info,
                    "scene": f"Scene {illustration.get('scene_number')}",
                    "style": illustration.get("image_provider"),
                    "emotional_tones": illustration.get("emotional_tones", []),
                    "intensity_score": illustration.get("intensity_score"),
                    "prompt": illustration.get("prompt"),
                    "text_excerpt": illustration.get("text_excerpt"),
                    "size": illustration.get("file_size"),
                    "width": illustration.get("width"),
                    "height": illustration.get("height"),
                    "created_at": created_at_str,
                    "generation_status": illustration.get("generation_status", "completed"),
                }
                images.append(image_data)

            # Get manuscript title (fallback to file-based system for now)
            manuscript_title = "Unknown Manuscript"
            try:
                manuscripts = get_saved_manuscripts()
                for manuscript in manuscripts:
                    generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
                    if generated_id == manuscript_id:
                        manuscript_title = manuscript.metadata.title
                        break
            except Exception:
                pass

            return {
                "images": images,
                "total_count": len(images),
                "manuscript_id": manuscript_id,
                "manuscript_title": manuscript_title
            }

        finally:
            illustration_service.close()

    except Exception as e:
        # Fallback to filesystem scanning if database fails
        print(f"Database query failed, falling back to filesystem: {e}")

        manuscripts = get_saved_manuscripts()

        # Find the manuscript
        manuscript_found = None
        for manuscript in manuscripts:
            generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
            if generated_id == manuscript_id:
                manuscript_found = manuscript
                break

        if not manuscript_found:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Manuscript not found"
            )

        # Fallback to filesystem scanning
        generated_images_dir = Path("illustrator_output") / "generated_images"
        images = []

        if generated_images_dir.exists():
            image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'}

            for image_file in generated_images_dir.iterdir():
                if image_file.is_file() and image_file.suffix.lower() in image_extensions:
                    filename = image_file.name
                    chapter_info = "Unknown"
                    scene_info = "Unknown"

                    if "chapter_" in filename and "scene_" in filename:
                        try:
                            parts = filename.split("_")
                            chapter_idx = next(i for i, part in enumerate(parts) if part == "chapter")
                            scene_idx = next(i for i, part in enumerate(parts) if part == "scene")

                            if chapter_idx + 1 < len(parts) and scene_idx + 1 < len(parts):
                                chapter_num = parts[chapter_idx + 1]
                                scene_num = parts[scene_idx + 1].split(".")[0]
                                chapter_info = f"Chapter {int(chapter_num)}"
                                scene_info = f"Scene {int(scene_num)}"
                        except (ValueError, IndexError, StopIteration):
                            pass

                    image_data = {
                        "url": f"/generated/{image_file.name}",
                        "filename": image_file.name,
                        "title": f"{chapter_info} - {scene_info}",
                        "description": f"Generated illustration for {chapter_info}, {scene_info}",
                        "chapter": chapter_info,
                        "scene": scene_info,
                        "style": "Generated illustration",
                        "size": image_file.stat().st_size,
                        "created_at": datetime.fromtimestamp(image_file.stat().st_mtime).isoformat()
                    }
                    images.append(image_data)

        # Sort images by chapter and scene
        def sort_key(img):
            try:
                filename = img["filename"]
                if "chapter_" in filename and "scene_" in filename:
                    parts = filename.split("_")
                    chapter_idx = next(i for i, part in enumerate(parts) if part == "chapter")
                    scene_idx = next(i for i, part in enumerate(parts) if part == "scene")

                    if chapter_idx + 1 < len(parts) and scene_idx + 1 < len(parts):
                        chapter_num = int(parts[chapter_idx + 1])
                        scene_num = int(parts[scene_idx + 1].split(".")[0])
                        return (chapter_num, scene_num)
            except (ValueError, IndexError, StopIteration):
                pass
            return (999, 999)

        images.sort(key=sort_key)

        return {
            "images": images,
            "total_count": len(images),
            "manuscript_id": manuscript_id,
            "manuscript_title": manuscript_found.metadata.title
        }


@router.delete("/{manuscript_id}/images/{image_id}")
async def delete_manuscript_image(manuscript_id: str, image_id: str) -> SuccessResponse:
    """Delete a specific image for a manuscript."""
    logger.debug(f"API CALL: DELETE /{{manuscript_id}}/images/{{image_id}} (delete_manuscript_image) with manuscript_id={manuscript_id}, image_id={image_id}")
    try:
        from illustrator.services.illustration_service import IllustrationService

        # Initialize illustration service
        illustration_service = IllustrationService()

        try:
            # Try to delete from database first
            illustration = illustration_service.get_illustration_by_id(image_id)

            if not illustration:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Image not found"
                )

            # Verify the image belongs to the correct manuscript
            if str(illustration.get("manuscript_id")) != manuscript_id:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Image does not belong to this manuscript"
                )

            # Delete the physical file
            file_path = illustration.get("file_path")
            if file_path and Path(file_path).exists():
                Path(file_path).unlink()

            # Delete from database
            illustration_service.delete_illustration(image_id)

            return SuccessResponse(
                message="Image deleted successfully"
            )

        finally:
            illustration_service.close()

    except HTTPException:
        raise
    except Exception as e:
        # Fallback to filesystem deletion if database fails
        print(f"Database delete failed, falling back to filesystem: {e}")

        # Parse image_id as filename for filesystem fallback
        filename = image_id
        if not filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
            filename = f"{image_id}.png"  # Assume PNG if no extension

        generated_images_dir = Path("illustrator_output") / "generated_images"
        image_path = generated_images_dir / filename

        if not image_path.exists():
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Image not found"
            )

        try:
            image_path.unlink()
            return SuccessResponse(
                message="Image deleted successfully"
            )
        except Exception as delete_error:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error deleting image: {str(delete_error)}"
            )


@router.delete("/{manuscript_id}/images")
async def delete_manuscript_images(manuscript_id: str) -> SuccessResponse:
    """Delete all images associated with a manuscript."""
    logger.debug(
        "API CALL: DELETE /{manuscript_id}/images (delete_manuscript_images) with manuscript_id=%s",
        manuscript_id,
    )

    db_removed = 0
    try:
        from illustrator.services.illustration_service import IllustrationService

        illustration_service = IllustrationService()
        try:
            db_removed = illustration_service.delete_illustrations_for_manuscript(manuscript_id)
        finally:
            illustration_service.close()
    except HTTPException:
        raise
    except Exception as exc:
        logger.warning(
            "Database image purge failed for manuscript %s: %s. Falling back to filesystem purge.",
            manuscript_id,
            exc,
        )

    if db_removed:
        return SuccessResponse(
            message="All images deleted successfully",
            data={"deleted_count": db_removed}
        )

    filesystem_removed = filesystem_delete_generated_images(manuscript_id)
    if filesystem_removed:
        return SuccessResponse(
            message="All images deleted successfully",
            data={"deleted_count": filesystem_removed, "source": "filesystem"}
        )

    return SuccessResponse(
        message="No images found for manuscript",
        data={"deleted_count": 0}
    )


@router.delete("/{manuscript_id}")
async def delete_manuscript(manuscript_id: str) -> SuccessResponse:
    """Delete a manuscript."""
    logger.debug(f"API CALL: DELETE /{{manuscript_id}} (delete_manuscript) with manuscript_id={manuscript_id}")
    manuscripts = get_saved_manuscripts()

    # Find and delete manuscript
    for manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, manuscript.file_path))
        if generated_id == manuscript_id:
            try:
                # Delete the file
                Path(manuscript.file_path).unlink(missing_ok=True)

                # Also delete generated images directory if it exists
                safe_title = manuscript.metadata.title.replace(" ", "_")
                images_dir = ILLUSTRATOR_OUTPUT_DIR / safe_title
                if images_dir.exists():
                    import shutil
                    shutil.rmtree(images_dir)

                # Invalidate cache to ensure the deleted manuscript doesn't appear in the dashboard
                invalidate_manuscripts_cache()

                return SuccessResponse(
                    message="Manuscript deleted successfully"
                )
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Error deleting manuscript: {str(e)}"
                )

    raise HTTPException(
        status_code=status.HTTP_404_NOT_FOUND,
        detail="Manuscript not found"
    )


@router.post("/{manuscript_id}/export")
async def export_manuscript(
    manuscript_id: str,
    export_format: str = "pdf"
) -> Dict[str, Any]:
    """Export manuscript in various formats (PDF, DOCX, HTML, JSON)."""
    logger.debug(f"API CALL: POST /{{manuscript_id}}/export (export_manuscript) with manuscript_id={manuscript_id}")
    # Find the manuscript
    manuscripts = get_saved_manuscripts()
    manuscript = None

    for saved_manuscript in manuscripts:
        generated_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, saved_manuscript.file_path))
        if generated_id == manuscript_id:
            manuscript = saved_manuscript
            break

    if not manuscript:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Manuscript not found"
        )

    # Validate export format
    supported_formats = ["pdf", "docx", "html", "json"]
    if export_format.lower() not in supported_formats:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported export format. Supported formats: {', '.join(supported_formats)}"
        )

    try:
        # Create exports directory
        exports_dir = ILLUSTRATOR_OUTPUT_DIR / "exports"
        exports_dir.mkdir(parents=True, exist_ok=True)

        # Generate filename
        safe_title = "".join(c for c in manuscript.metadata.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        export_format_lower = export_format.lower()

        if export_format_lower == "json":
            # JSON Export - raw data
            filename = f"{safe_title}_{timestamp}.json"
            file_path = exports_dir / filename

            export_data = {
                "manuscript_id": manuscript_id,
                "metadata": manuscript.metadata.model_dump(),
                "chapters": [chapter.model_dump() for chapter in manuscript.chapters],
                "export_info": {
                    "format": "json",
                    "exported_at": datetime.now().isoformat(),
                    "total_chapters": len(manuscript.chapters),
                    "total_words": sum(len(chapter.content.split()) for chapter in manuscript.chapters)
                }
            }

            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)

        elif export_format_lower == "html":
            # HTML Export
            filename = f"{safe_title}_{timestamp}.html"
            file_path = exports_dir / filename

            html_content = generate_html_export(manuscript)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

        elif export_format_lower == "pdf":
            # PDF Export
            filename = f"{safe_title}_{timestamp}.pdf"
            file_path = exports_dir / filename

            await generate_pdf_export(manuscript, file_path)

        elif export_format_lower == "docx":
            # DOCX Export
            filename = f"{safe_title}_{timestamp}.docx"
            file_path = exports_dir / filename

            generate_docx_export(manuscript, file_path)

        # Return download info
        return {
            "success": True,
            "export_format": export_format.upper(),
            "filename": filename,
            "download_url": f"/api/manuscripts/{manuscript_id}/download/{filename}",
            "file_size": file_path.stat().st_size,
            "exported_at": datetime.now().isoformat(),
            "manuscript_title": manuscript.metadata.title
        }

    except Exception as e:
        import traceback
        print(f"Export error: {traceback.format_exc()}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to export manuscript: {str(e)}"
        )


def generate_html_export(manuscript: SavedManuscript) -> str:
    """Generate HTML export of the manuscript."""
    html_template = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{title} by {author}</title>
        <style>
            body {{
                font-family: 'Times New Roman', serif;
                line-height: 1.6;
                max-width: 800px;
                margin: 0 auto;
                padding: 2rem;
                background-color: #f8f9fa;
            }}
            .manuscript-header {{
                text-align: center;
                margin-bottom: 3rem;
                padding: 2rem;
                background: white;
                border-radius: 8px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }}
            .manuscript-title {{
                font-size: 2.5rem;
                margin-bottom: 0.5rem;
                color: #2c3e50;
            }}
            .manuscript-author {{
                font-size: 1.2rem;
                color: #6c757d;
                margin-bottom: 1rem;
            }}
            .manuscript-meta {{
                font-size: 0.9rem;
                color: #868e96;
            }}
            .chapter {{
                background: white;
                margin: 2rem 0;
                padding: 2rem;
                border-radius: 8px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }}
            .chapter-title {{
                font-size: 1.8rem;
                margin-bottom: 1.5rem;
                color: #2c3e50;
                border-bottom: 2px solid #3498db;
                padding-bottom: 0.5rem;
            }}
            .chapter-content {{
                text-align: justify;
                text-indent: 2rem;
            }}
            .chapter-content p {{
                margin-bottom: 1rem;
            }}
            @media print {{
                body {{
                    background-color: white;
                    max-width: none;
                }}
                .manuscript-header, .chapter {{
                    box-shadow: none;
                    border: 1px solid #ddd;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="manuscript-header">
            <h1 class="manuscript-title">{title}</h1>
            <p class="manuscript-author">by {author}</p>
            <div class="manuscript-meta">
                <p>Genre: {genre} | Chapters: {chapter_count} | Words: {word_count:,}</p>
                <p>Exported: {export_date}</p>
            </div>
        </div>

        {chapters_html}
    </body>
    </html>
    """

    # Generate chapters HTML
    chapters_html = ""
    for chapter in manuscript.chapters:
        # Convert line breaks to paragraphs
        content_paragraphs = [p.strip() for p in chapter.content.split('\n\n') if p.strip()]
        content_html = '\n'.join([f'<p>{p}</p>' for p in content_paragraphs])

        chapters_html += f"""
        <div class="chapter">
            <h2 class="chapter-title">Chapter {chapter.number}: {chapter.title}</h2>
            <div class="chapter-content">
                {content_html}
            </div>
        </div>
        """

    return html_template.format(
        title=manuscript.metadata.title,
        author=manuscript.metadata.author or "Unknown Author",
        genre=manuscript.metadata.genre or "Fiction",
        chapter_count=len(manuscript.chapters),
        word_count=sum(len(chapter.content.split()) for chapter in manuscript.chapters),
        export_date=datetime.now().strftime("%B %d, %Y"),
        chapters_html=chapters_html
    )


async def generate_pdf_export(manuscript: SavedManuscript, file_path: Path):
    """Generate PDF export using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    except ImportError:
        # Fallback: generate HTML and convert to PDF using weasyprint if available
        try:
            import weasyprint
            html_content = generate_html_export(manuscript)
            weasyprint.HTML(string=html_content).write_pdf(str(file_path))
            return
        except ImportError:
            raise Exception("PDF generation requires 'reportlab' or 'weasyprint' package. Install with: pip install reportlab")

    # Create PDF document
    doc = SimpleDocTemplate(str(file_path), pagesize=letter, topMargin=1*inch)
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER
    )

    author_style = ParagraphStyle(
        'CustomAuthor',
        parent=styles['Normal'],
        fontSize=14,
        spaceAfter=20,
        alignment=TA_CENTER
    )

    chapter_title_style = ParagraphStyle(
        'ChapterTitle',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=20,
        spaceBefore=30
    )

    content_style = ParagraphStyle(
        'Content',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        firstLineIndent=20
    )

    # Title page
    story.append(Paragraph(manuscript.metadata.title, title_style))
    story.append(Paragraph(f"by {manuscript.metadata.author or 'Unknown Author'}", author_style))
    story.append(Spacer(1, 0.5*inch))

    # Metadata
    metadata_text = f"""
    Genre: {manuscript.metadata.genre or 'Fiction'}<br/>
    Chapters: {len(manuscript.chapters)}<br/>
    Words: {sum(len(chapter.content.split()) for chapter in manuscript.chapters):,}<br/>
    Exported: {datetime.now().strftime('%B %d, %Y')}
    """
    story.append(Paragraph(metadata_text, styles['Normal']))
    story.append(PageBreak())

    # Chapters
    for i, chapter in enumerate(manuscript.chapters):
        if i > 0:
            story.append(PageBreak())

        # Chapter title
        story.append(Paragraph(f"Chapter {chapter.number}: {chapter.title}", chapter_title_style))

        # Chapter content
        paragraphs = [p.strip() for p in chapter.content.split('\n\n') if p.strip()]
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, content_style))

    # Build PDF
    doc.build(story)


def generate_docx_export(manuscript: SavedManuscript, file_path: Path):
    """Generate DOCX export using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise Exception("DOCX generation requires 'python-docx' package. Install with: pip install python-docx")

    # Create document
    doc = Document()

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(manuscript.metadata.title)
    title_run.bold = True
    title_run.font.size = doc.styles['Title'].font.size

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(f"by {manuscript.metadata.author or 'Unknown Author'}")
    author_run.italic = True

    doc.add_paragraph()  # Spacer

    # Metadata
    metadata_para = doc.add_paragraph()
    metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_text = f"""Genre: {manuscript.metadata.genre or 'Fiction'}
Chapters: {len(manuscript.chapters)}
Words: {sum(len(chapter.content.split()) for chapter in manuscript.chapters):,}
Exported: {datetime.now().strftime('%B %d, %Y')}"""
    metadata_para.add_run(metadata_text)

    doc.add_page_break()

    # Chapters
    for chapter in manuscript.chapters:
        # Chapter title
        chapter_heading = doc.add_heading(f"Chapter {chapter.number}: {chapter.title}", level=1)

        # Chapter content
        paragraphs = [p.strip() for p in chapter.content.split('\n\n') if p.strip()]
        for paragraph in paragraphs:
            para = doc.add_paragraph(paragraph)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

    # Remove last page break
    if doc.paragraphs:
        last_para = doc.paragraphs[-1]
        if last_para._p.getparent() is not None:
            last_para._p.getparent().remove(last_para._p)

    # Save document
    doc.save(str(file_path))
